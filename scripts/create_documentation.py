"""
Proje Dokümantasyonunu DOCX formatında oluşturur.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def create_documentation():
    """Profesyonel DOCX dokümantasyonu oluştur."""
    doc = Document()
    
    # Stil ayarları
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # ========== KAPAK SAYFASI ==========
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title_run = title.add_run('🤖 AI DECISION ECOSYSTEM ENGINE')
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0, 102, 204)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Multi-Agent Karar Destek Sistemi')
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.add_run('Proje Dokümantasyonu').bold = True
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph()
    date_para.add_run('Mart 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========== İÇİNDEKİLER ==========
    doc.add_heading('İÇİNDEKİLER', level=1)
    
    toc_items = [
        ('1. Proje Nedir?', '3'),
        ('2. Problem ve Çözüm', '4'),
        ('3. Sistem Nasıl Çalışır?', '5'),
        ('4. Agent\'lar Kimdir?', '7'),
        ('5. Senaryo Sınıflandırma', '10'),
        ('6. Karar Süreci', '12'),
        ('7. Gerçek Hayat Örneği', '14'),
        ('8. Mevcut Durum', '16'),
        ('9. Gelecek Planları', '17'),
        ('10. Sözlük', '18'),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{item}').bold = True
        p.add_run(f' {"." * (50 - len(item))} {page}')
    
    doc.add_page_break()
    
    # ========== BÖLÜM 1: PROJE NEDİR? ==========
    doc.add_heading('1. PROJE NEDİR?', level=1)
    
    doc.add_heading('Tek Cümleyle', level=2)
    quote = doc.add_paragraph()
    quote_run = quote.add_run('Şirketlerin önemli kararlarını 3 farklı uzman yapay zeka ile değerlendiren bir karar destek sistemi.')
    quote_run.italic = True
    quote_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_heading('Detaylı Açıklama', level=2)
    doc.add_paragraph(
        'Bir şirket büyük bir karar vermek istediğinde (örneğin: "25 milyon dolarlık yapay zeka '
        'yatırımı yapalım mı?"), genellikle farklı departmanların görüşü alınır:'
    )
    
    bullets = [
        'CEO stratejik açıdan bakar',
        'CFO finansal açıdan bakar',
        'HR insan kaynakları açısından bakar'
    ]
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_paragraph(
        'Bu proje, bu süreci otomatikleştiren bir sistemdir. 3 yapay zeka ajanı, her biri kendi '
        'uzmanlık alanından senaryoyu değerlendirir ve ortak bir karar üretir.'
    )
    
    # Akış şeması tablosu
    doc.add_heading('Sistem Akışı', level=2)
    
    flow_table = doc.add_table(rows=5, cols=3)
    flow_table.style = 'Table Grid'
    
    flow_data = [
        ['Girdi', '→', 'Değerlendirme'],
        ['Senaryo Bilgileri', '→', '3 Agent Analizi'],
        ['(Bütçe, ROI, Risk)', '→', '(CEO, CFO, HR)'],
        ['', '↓', ''],
        ['', 'KARAR', ''],
    ]
    
    for i, row_data in enumerate(flow_data):
        row = flow_table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========== BÖLÜM 2: PROBLEM VE ÇÖZÜM ==========
    doc.add_heading('2. PROBLEM VE ÇÖZÜM', level=1)
    
    doc.add_heading('Problem: Tek Bakış Açısı Yetersiz', level=2)
    doc.add_paragraph('Geleneksel karar verme süreçlerinde:')
    
    problem_table = doc.add_table(rows=4, cols=2)
    problem_table.style = 'Table Grid'
    
    problem_data = [
        ['Durum', 'Sonuç'],
        ['Sadece CEO karar verirse', 'Finansal riskler göz ardı edilebilir'],
        ['Sadece CFO karar verirse', 'Stratejik fırsatlar kaçırılabilir'],
        ['Sadece HR karar verirse', 'Büyüme hedefleri aksayabilir'],
    ]
    
    for i, row_data in enumerate(problem_data):
        row = problem_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_heading('Çözüm: Çoklu Perspektif', level=2)
    doc.add_paragraph(
        'Sistemimiz 3 farklı bakış açısını bir araya getirir. Her agent kendi uzmanlık '
        'alanından değerlendirme yapar ve sonuçlar ağırlıklı olarak birleştirilir.'
    )
    
    solution_table = doc.add_table(rows=3, cols=2)
    solution_table.style = 'Table Grid'
    
    solution_data = [
        ['Geleneksel', 'Bizim Sistem'],
        ['Tek perspektif → Riskli', 'Çoklu perspektif → Dengeli'],
        ['Hızlı ama hatalı', 'Kapsamlı ve güvenilir'],
    ]
    
    for i, row_data in enumerate(solution_data):
        row = solution_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # ========== BÖLÜM 3: SİSTEM NASIL ÇALIŞIR? ==========
    doc.add_heading('3. SİSTEM NASIL ÇALIŞIR?', level=1)
    
    doc.add_heading('Ana Akış (5 Adım)', level=2)
    
    steps = [
        ('ADIM 1: Senaryo Girişi', 'Kullanıcı 4 temel bilgi girer: Bütçe, ROI, Risk, Ekip Hazırlığı'),
        ('ADIM 2: Sınıflandırma', 'Sistem senaryonun tipini belirler (5 tipten biri)'),
        ('ADIM 3: Agent Analizi', 'Her agent (CEO, CFO, HR) kendi perspektifinden değerlendirir'),
        ('ADIM 4: Ağırlıklı Hesaplama', 'Senaryo tipine göre ağırlıklar uygulanır'),
        ('ADIM 5: Final Karar', 'APPROVE / REVISE / REJECT kararı üretilir'),
    ]
    
    for step_title, step_desc in steps:
        p = doc.add_paragraph()
        p.add_run(step_title).bold = True
        doc.add_paragraph(step_desc)
    
    doc.add_heading('Senaryo Giriş Parametreleri', level=2)
    
    param_table = doc.add_table(rows=5, cols=3)
    param_table.style = 'Table Grid'
    
    param_data = [
        ['Alan', 'Açıklama', 'Örnek'],
        ['Bütçe', 'Proje maliyeti (milyon $)', '25'],
        ['ROI', 'Beklenen getiri (%)', '45'],
        ['Risk', 'Risk seviyesi (1-10)', '5'],
        ['Ekip Hazırlığı', 'Mevcut ekip yeterliliği (1-10)', '7'],
    ]
    
    for i, row_data in enumerate(param_data):
        row = param_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # ========== BÖLÜM 4: AGENT'LAR KİMDİR? ==========
    doc.add_heading('4. AGENT\'LAR KİMDİR?', level=1)
    
    agents = [
        {
            'name': 'CEO AGENT - Strateji Uzmanı 🎯',
            'role': 'Chief Executive Officer (Genel Müdür)',
            'focus': ['Şirket stratejisi', 'Büyüme hedefleri', 'Pazar fırsatları', 'Rekabet avantajı', 'Uzun vadeli vizyon'],
            'question': 'Bu proje şirketi ileriye taşır mı?',
            'example': 'ROI %45 ile güçlü bir getiri. Stratejik hedeflerimizle uyumlu. DESTEK veriyorum.'
        },
        {
            'name': 'CFO AGENT - Finans Uzmanı 💰',
            'role': 'Chief Financial Officer (Finans Direktörü)',
            'focus': ['Bütçe yönetimi', 'Maliyet analizi', 'Nakit akışı', 'Yatırım getirisi', 'Finansal risk'],
            'question': 'Bu projenin finansal mantığı var mı?',
            'example': '25M$ yatırım için %45 ROI kabul edilebilir. Risk seviyesi orta. FİNANSAL OLARAK UYGUN.'
        },
        {
            'name': 'HR AGENT - İnsan Kaynakları Uzmanı 👥',
            'role': 'Human Resources (İnsan Kaynakları)',
            'focus': ['Ekip kapasitesi', 'Yetenek ihtiyacı', 'İşe alım gereksinimleri', 'Eğitim ihtiyaçları', 'Organizasyonel etki'],
            'question': 'Bu projeyi yapacak ekibimiz var mı?',
            'example': 'Ekip hazırlığı 3/10. Bu proje için en az 8 yeni çalışan gerekli. NÖTR - Önce işe alım.'
        },
    ]
    
    for agent in agents:
        doc.add_heading(agent['name'], level=2)
        
        p = doc.add_paragraph()
        p.add_run('Rol: ').bold = True
        p.add_run(agent['role'])
        
        p = doc.add_paragraph()
        p.add_run('Odak Alanları:').bold = True
        for focus in agent['focus']:
            doc.add_paragraph(focus, style='List Bullet')
        
        p = doc.add_paragraph()
        p.add_run('Temel Soru: ').bold = True
        p.add_run(f'"{agent["question"]}"').italic = True
        
        p = doc.add_paragraph()
        p.add_run('Örnek Çıktı: ').bold = True
        p.add_run(f'"{agent["example"]}"').italic = True
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ========== BÖLÜM 5: SENARYO SINIFLANDIRMA ==========
    doc.add_heading('5. SENARYO SINIFLANDIRMA', level=1)
    
    doc.add_paragraph('Sistem, her senaryoyu otomatik olarak 5 tipten birine sınıflandırır:')
    
    types_table = doc.add_table(rows=6, cols=4)
    types_table.style = 'Table Grid'
    
    types_data = [
        ['Tip', 'Özellikler', 'Örnek', 'Baskın Agent'],
        ['HIGH_GROWTH', 'Yüksek ROI (>%30), Düşük risk', 'Yeni pazara giriş', 'CEO (%40)'],
        ['COST_OPTIMIZATION', 'Düşük bütçe (<$10M), Tasarruf odaklı', 'Süreç iyileştirme', 'CFO (%50)'],
        ['TEAM_EXPANSION', 'Düşük ekip hazırlığı (<5)', 'Yeni departman', 'HR (%50)'],
        ['STRATEGIC_PIVOT', 'Yüksek risk (>7), Strateji değişikliği', 'İş modeli değişikliği', 'CEO (%45)'],
        ['MAINTENANCE', 'Düşük her şey, Rutin işler', 'Sistem güncellemesi', 'Eşit (%33)'],
    ]
    
    for i, row_data in enumerate(types_data):
        row = types_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_heading('Ağırlık Dağılımı', level=2)
    
    weights_table = doc.add_table(rows=6, cols=4)
    weights_table.style = 'Table Grid'
    
    weights_data = [
        ['Senaryo Tipi', 'CEO', 'CFO', 'HR'],
        ['high_growth', '40%', '35%', '25%'],
        ['cost_optimization', '25%', '50%', '25%'],
        ['team_expansion', '25%', '25%', '50%'],
        ['strategic_pivot', '45%', '30%', '25%'],
        ['maintenance', '33%', '33%', '33%'],
    ]
    
    for i, row_data in enumerate(weights_data):
        row = weights_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # ========== BÖLÜM 6: KARAR SÜRECİ ==========
    doc.add_heading('6. KARAR SÜRECİ', level=1)
    
    doc.add_heading('Karar Eşikleri', level=2)
    
    threshold_table = doc.add_table(rows=4, cols=3)
    threshold_table.style = 'Table Grid'
    
    threshold_data = [
        ['Skor Aralığı', 'Karar', 'Anlam'],
        ['0-49', 'REJECT ❌', 'Projeyi reddet'],
        ['50-69', 'REVISE ⚠️', 'Revizyonla tekrar değerlendir'],
        ['70-100', 'APPROVE ✅', 'Projeyi onayla'],
    ]
    
    for i, row_data in enumerate(threshold_data):
        row = threshold_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_heading('Hesaplama Formülü', level=2)
    
    formula = doc.add_paragraph()
    formula.add_run('Final Skor = ').bold = True
    formula.add_run('(CEO Skoru × CEO Ağırlığı) + (CFO Skoru × CFO Ağırlığı) + (HR Skoru × HR Ağırlığı)')
    
    doc.add_heading('Hesaplama Örneği', level=2)
    
    example_data = [
        'Senaryo: AI Platform Yatırımı',
        'Bütçe: $25M, ROI: %45, Risk: 5, Ekip: 3',
        'Tip: team_expansion (ekip hazırlığı düşük)',
        '',
        'Agent Skorları:',
        '• CEO: 85 (Strateji uygun)',
        '• CFO: 90 (Finans olumlu)',
        '• HR: 50 (Ekip yetersiz)',
        '',
        'Ağırlıklar (team_expansion):',
        '• CEO: %25, CFO: %25, HR: %50',
        '',
        'Hesaplama:',
        '(85 × 0.25) + (90 × 0.25) + (50 × 0.50)',
        '= 21.25 + 22.50 + 25.00',
        '= 68.75',
        '',
        'Karar: REVISE (68.75 < 70)',
        'Yorum: Ekip eksikliğini giderin, sonra onaylayın.',
    ]
    
    for line in example_data:
        if line.startswith('•'):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line:
            doc.add_paragraph(line)
    
    doc.add_page_break()
    
    # ========== BÖLÜM 7: GERÇEK HAYAT ÖRNEĞİ ==========
    doc.add_heading('7. GERÇEK HAYAT ÖRNEĞİ', level=1)
    
    doc.add_heading('Senaryo: E-Ticaret Şirketi AI Yatırımı', level=2)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Şirket: ').bold = True
    p.add_run('TechMart E-Ticaret')
    
    p = doc.add_paragraph()
    p.add_run('Karar: ').bold = True
    p.add_run('Yapay zeka tabanlı öneri sistemi yatırımı yapılsın mı?')
    
    doc.add_paragraph()
    doc.add_heading('Proje Bilgileri:', level=3)
    
    project_info = [
        'Bütçe: $15 milyon',
        'Beklenen ROI: %60',
        'Risk Seviyesi: 4/10',
        'Mevcut Ekip Hazırlığı: 8/10',
    ]
    for info in project_info:
        doc.add_paragraph(info, style='List Bullet')
    
    doc.add_heading('Sonuç:', level=3)
    
    results = [
        'Sınıflandırma: HIGH_GROWTH (ROI %60 yüksek)',
        'CEO Skoru: 92 ("Stratejik olarak destekliyorum")',
        'CFO Skoru: 88 ("Finansal olarak uygun")',
        'HR Skoru: 85 ("Ekip hazır, 2 ek mühendis yeterli")',
        '',
        'Hesaplama: (92×0.40) + (88×0.35) + (85×0.25) = 88.85',
        '',
        'Final Karar: APPROVE ✅',
        'Öneri: Proje onaylanabilir. 2 ek mühendis işe alımı yapılmalı.',
    ]
    
    for result in results:
        if result:
            doc.add_paragraph(result)
    
    doc.add_page_break()
    
    # ========== BÖLÜM 8: MEVCUT DURUM ==========
    doc.add_heading('8. MEVCUT DURUM', level=1)
    
    doc.add_heading('Tamamlanan Özellikler', level=2)
    
    completed_table = doc.add_table(rows=11, cols=2)
    completed_table.style = 'Table Grid'
    
    completed_data = [
        ['#', 'Özellik'],
        ['1', 'Multi-Agent Sistemi (CEO, CFO, HR)'],
        ['2', 'Round-Based Tartışma'],
        ['3', 'Konsensüs Algılama'],
        ['4', 'Senaryo Sınıflandırma (5 tip)'],
        ['5', 'Dinamik Ağırlıklar'],
        ['6', 'REST API (8 endpoint)'],
        ['7', 'PostgreSQL Veritabanı'],
        ['8', 'Docker Deployment'],
        ['9', 'HTML Dashboard'],
        ['10', '86 Test + 105 Senaryo'],
    ]
    
    for i, row_data in enumerate(completed_data):
        row = completed_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_heading('Mevcut Kısıtlamalar', level=2)
    
    limitations = [
        'Sınıflandırma Kural Tabanlı: ML yerine IF/ELSE kuralları kullanılıyor',
        'Agentlar Tepki Vermiyor: Birbirlerini görüyor ama fikir değiştirmiyor',
        'Veri Öğrenimde Kullanılmıyor: 105 senaryo sadece depolanıyor',
    ]
    for limit in limitations:
        doc.add_paragraph(limit, style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Gerçek Hayata Yakınlık: ').bold = True
    p.add_run('%70-75')
    
    doc.add_page_break()
    
    # ========== BÖLÜM 9: GELECEK PLANLARI ==========
    doc.add_heading('9. GELECEK PLANLARI', level=1)
    
    future_table = doc.add_table(rows=5, cols=3)
    future_table.style = 'Table Grid'
    
    future_data = [
        ['Faz', 'Özellik', 'Açıklama'],
        ['8', 'Gerçek ML Sınıflandırma', 'scikit-learn ile model eğitimi'],
        ['9', 'Reaktif Agentlar', 'Agentlar birbirinin skoruna tepki versin'],
        ['10', 'Feedback Loop', 'Model performans ölçümü ve yeniden eğitim'],
        ['11', 'LLM Entegrasyonu', 'OpenAI/Ollama ile doğal dil desteği'],
    ]
    
    for i, row_data in enumerate(future_data):
        row = future_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # ========== BÖLÜM 10: SÖZLÜK ==========
    doc.add_heading('10. SÖZLÜK', level=1)
    
    glossary_table = doc.add_table(rows=17, cols=2)
    glossary_table.style = 'Table Grid'
    
    glossary_data = [
        ['Terim', 'Açıklama'],
        ['Agent', 'Belirli bir rolde çalışan yapay zeka modülü'],
        ['Orkestratör', 'Agent\'ları koordine eden sistem'],
        ['Senaryo', 'Değerlendirilecek iş kararı'],
        ['Sınıflandırma', 'Senaryonun tipini belirleme'],
        ['Ağırlık', 'Agent\'ın karardaki etkisi (%)'],
        ['ROI', 'Return on Investment - Yatırım getirisi'],
        ['Round', 'Agent\'ların görüş bildirdiği tur'],
        ['Konsensüs', 'Tüm agent\'ların benzer görüşte olması'],
        ['Agregasyon', 'Skorların birleştirilmesi'],
        ['APPROVE', 'Projenin onaylanması'],
        ['REVISE', 'Revizyona gönderilmesi'],
        ['REJECT', 'Projenin reddedilmesi'],
        ['ML', 'Machine Learning - Makine Öğrenimi'],
        ['LLM', 'Large Language Model - Büyük Dil Modeli'],
        ['Kural Tabanlı', 'IF/ELSE kurallarıyla çalışan sistem'],
        ['Reaktif', 'Diğer görüşlere tepki veren'],
    ]
    
    for i, row_data in enumerate(glossary_data):
        row = glossary_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.add_run('Dokümantasyon Tarihi: ').bold = True
    footer.add_run('Mart 2026')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Kaydet
    output_path = Path(__file__).parent.parent / 'docs' / 'PROJE_DOKUMANTASYONU.docx'
    doc.save(str(output_path))
    print(f'✅ DOCX dosyası oluşturuldu: {output_path}')
    
    return str(output_path)


if __name__ == '__main__':
    create_documentation()
